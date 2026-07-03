from __future__ import annotations

import json
import textwrap
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import build_pdf_report

ROOT = Path(__file__).resolve().parent
EVIDENCE = ROOT / "evidence"
SCREENSHOTS = ROOT / "screenshots"
FINAL_DIR = ROOT.parent / "final_submission"
DOCX_PATH = FINAL_DIR / "GXXX_SEML_Assignment_01_Ecommerce_Recommendation_Final_Report.docx"

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
INK = RGBColor(0x1F, 0x29, 0x33)
MUTED = RGBColor(0x52, 0x60, 0x6D)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, margin_value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_in: float):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[float]):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10

    code_style = styles.add_style("Code Block", 1)
    code_style.font.name = "Consolas"
    code_style._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code_style.font.size = Pt(8)
    code_style.paragraph_format.space_after = Pt(3)
    code_style.paragraph_format.line_spacing = 1.0

    header = section.header.paragraphs[0]
    header.text = ""
    r1 = header.add_run("AIMLCZG546 - Software Engineering for Machine Learning")
    set_run_font(r1, size=9, color=MUTED)
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT

    footer = section.footer.paragraphs[0]
    footer.text = ""
    add_page_number(footer)


def add_title_block(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("AIMLCZG546 - Software Engineering for Machine Learning")
    set_run_font(r, size=18, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Assignment I: Real-Time Product Recommendation for an E-commerce Platform")
    set_run_font(r, size=17, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("Domain: E-commerce / Retail | Patterns: Event-Driven Architecture and API Gateway")
    set_run_font(r, size=11, color=MUTED)


def add_para(doc: Document, text: str, style: str | None = None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run_font(r, size=11, color=INK)
    return p


def add_callout(doc: Document, text: str, fill="EAF3FB"):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)
    doc.add_paragraph()


def add_data_table(doc: Document, rows: list[list[str]], widths: list[float], header=True):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for row_idx, row in enumerate(rows):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            set_cell_margins(cell)
            if header and row_idx == 0:
                set_cell_shading(cell, "F2F4F7")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(value))
            set_run_font(
                run,
                size=9,
                color=DARK_BLUE if header and row_idx == 0 else INK,
                bold=bool(header and row_idx == 0),
            )
    doc.add_paragraph()
    return table


def add_image(doc: Document, path: Path, caption: str, width=6.3):
    if not path.exists():
        return
    doc.add_picture(str(path), width=Inches(width))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    set_run_font(r, size=9, color=MUTED, italic=True)


def add_code_block(doc: Document, text: str):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F6F8FA")
    p = cell.paragraphs[0]
    p.style = doc.styles["Code Block"]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(textwrap.dedent(text).strip())
    set_run_font(r, name="Consolas", size=8, color=RGBColor(0x11, 0x18, 0x27))
    doc.add_paragraph()


def build_docx() -> Path:
    FINAL_DIR.mkdir(exist_ok=True)
    EVIDENCE.mkdir(exist_ok=True)
    SCREENSHOTS.mkdir(exist_ok=True)

    diagram_paths = build_pdf_report.generate_diagrams()
    api_evidence = build_pdf_report.make_api_evidence_image()
    consumer_log = build_pdf_report.make_consumer_log_image()
    gateway_swagger = SCREENSHOTS / "gateway_swagger_ui.png"
    service_swagger = SCREENSHOTS / "recommendation_service_swagger_ui.png"
    recommendation_plot = EVIDENCE / "recommendation_output_plot.png"
    sample_output = ""
    sample_output_path = EVIDENCE / "sample_output.txt"
    if sample_output_path.exists():
        sample_output = sample_output_path.read_text(encoding="utf-8", errors="replace")

    doc = Document()
    configure_document(doc)
    add_title_block(doc)
    add_callout(
        doc,
        "This editable report completes Objective 1 and Objective 2: GR4ML requirements formulation, "
        "architecture with ML and non-ML components, two architectural patterns, and a working prototype with run evidence.",
    )

    add_data_table(doc, [
        ["Field", "Value"],
        ["Group No", "GXXX - replace with actual group number before portal upload"],
        ["Submission file", "<groupid>.docx or <groupid>.pdf as required by Taxila"],
        ["Implementation notebook", "final_submission/GXXX.ipynb - rename GXXX to actual group number"],
        ["Submission deadline", "30 June 2026, 23:00"],
        ["Prototype folder", "seml-ecommerce-reco/"],
    ], [1.55, 4.95])
    add_data_table(doc, [
        ["Sl.", "BITS ID", "Name", "Contribution (qualitative)", "%"],
        ["1", "To fill", "To fill", "Requirements Lead: problem statement, requirements, measurable goals, quality requirements", "25"],
        ["2", "To fill", "To fill", "GR4ML Modelling Lead: Business, Analytics Design, and Data Preparation views", "25"],
        ["3", "To fill", "To fill", "Architecture Lead: architecture diagram and pattern selection / justification", "25"],
        ["4", "To fill", "To fill", "Implementation Lead: code, run evidence, screenshots, and code explanation", "25"],
    ], [0.35, 0.85, 0.95, 3.85, 0.5])

    doc.add_heading("Assignment Coverage Matrix", level=1)
    add_data_table(doc, [
        ["Official requirement", "Where covered"],
        ["Select a domain and define an ML-based problem statement", "Section 1.1"],
        ["Formulate requirement specifications and measurable goals using GR4ML concepts", "Sections 1.2 and 1.3"],
        ["Develop GR4ML Business View, Analytics Design View, and Data Preparation View", "Section 1.3 and Figures 1-3"],
        ["Identify the top three quality requirements with justification", "Section 1.4"],
        ["Draw architecture diagram showing ML and non-ML components", "Section 2.1 and Figure 4"],
        ["Select and apply any two relevant architectural patterns", "Sections 2.2 and 2.3"],
        ["Implement selected patterns using appropriate technologies", "Section 2.3, Appendix A, prototype folder, and GXXX.ipynb"],
        ["Include screenshots / outputs of the application with explanation", "Section 2.4 and Figures 5-9"],
        ["Mention group details and contribution", "Cover page and Task Distribution section"],
    ], [3.25, 3.25])

    doc.add_heading("Objective 1 - Requirements Formulation", level=1)
    doc.add_heading("1.1 Domain and Problem Statement", level=2)
    add_para(doc, "We select the e-commerce / online retail domain. A modern storefront earns incremental revenue through cross-sell and up-sell: showing each shopper additional products they are likely to want. The business goal is to lift revenue per session by replacing a static best-seller block with a personalised list that reacts to the user's latest behaviour.")
    add_para(doc, "Problem statement: given the evolving history of user-item interactions on the platform, build a system that ranks catalogue items for a user and serves a personalised top-N recommendation list in real time on the storefront. Formally, this is a recommendation / ranking task using implicit feedback such as views, clicks, add-to-cart actions, and purchases.")
    add_para(doc, "Why ML is appropriate: the relationship between past behaviour and future intent is sparse, high-dimensional, and constantly drifting as products, seasons, and user preferences change. Hand-written rules do not scale across a large catalogue; a data-driven model can learn item affinity from interaction patterns.")

    doc.add_heading("1.2 Requirement Specifications and Measurable Goals", level=2)
    add_data_table(doc, [
        ["ID", "Requirement", "GR4ML role", "Verification"],
        ["FR1", "Capture user activity events: view, click, cart, purchase, with user and item identifiers.", "Operational capability", "POST sample events and check they enter the queue."],
        ["FR2", "Continuously update user-item interaction features from the event stream.", "Analytics capability", "Consumer log shows processed events and feature count changes."],
        ["FR3", "Serve personalised top-N recommendations for any valid user, excluding already-seen items.", "ML output / insight", "GET /recommend returns ranked unseen items."],
        ["FR4", "Expose client traffic through one authenticated gateway.", "Software interface requirement", "External calls use gateway port 8000, not internal port 8001."],
    ], [0.45, 2.85, 1.45, 1.75])
    add_data_table(doc, [
        ["Quality / measurable goal", "Metric", "Target", "Reason"],
        ["Relevant top-N", "Offline Precision@5", ">= 0.30", "Measures whether top returned products include held-out relevant items."],
        ["Business uplift", "CTR uplift vs non-personalised block", ">= 8%", "Represents expected production A/B test improvement."],
        ["Real-time serving", "Recommendation API p95 latency", "< 150 ms", "Recommendations appear on the page-render path."],
        ["Freshness", "Event-to-feature lag", "< 60 s", "Recent behaviour should affect recommendations quickly."],
        ["Catalogue reach", "Catalogue coverage", ">= 60%", "Avoids showing only a narrow set of popular products."],
    ], [1.55, 1.35, 0.9, 2.7])
    add_callout(doc, "Executed prototype note: the local implementation reports offline Precision@5 = 0.323 on deterministic synthetic interaction data, clearing the >= 0.30 target.")

    doc.add_heading("1.3 GR4ML Views", level=2)
    add_para(doc, "GR4ML links business intent, analytics design, and data preparation. In this solution the traceable chain is: revenue and engagement goal -> product-display decision -> question about user intent -> top-N recommendation insight -> prediction/ranking analytics goal -> prepared interaction matrix and item-similarity table.")
    add_image(doc, diagram_paths[0], "Figure 1: Business View showing the strategic goal, decision, question, indicators, and insight.")
    add_image(doc, diagram_paths[1], "Figure 2: Analytics Design View showing the prediction goal, candidate algorithms, and softgoal trade-offs.")
    add_image(doc, diagram_paths[2], "Figure 3: Data Preparation View showing how raw events become prepared ranking features.")
    add_para(doc, "Business View explanation: the strategic goal is to increase revenue per session through cross-sell, measured by click-through rate and average order value. The operational decision is which products to show each user. The analytics question is what this user will most likely engage with, and the insight is a personalised top-N list.")
    add_para(doc, "Analytics Design View explanation: the insight becomes a prediction goal: rank items by user preference. Item-based collaborative filtering is selected for the real-time path because similarities can be precomputed and serving is fast. Matrix factorisation is retained as a future offline quality enhancement, while content-based filtering helps with cold-start items.")
    add_para(doc, "Data Preparation View explanation: user events, orders, and catalogue metadata are validated, deduplicated, weighted, and transformed into a user-item interaction matrix. Cosine similarity over item columns produces an item-similarity table; already-seen items are filtered before recommendations are returned.")

    doc.add_heading("1.4 Top Three Quality Requirements", level=2)
    add_data_table(doc, [
        ["Quality requirement", "Justification", "Design decision influenced"],
        ["Low latency", "Recommendations render inline while the shopper browses; slow serving damages the shopping experience.", "Precompute item similarity, cache top-N, keep the serving path simple."],
        ["Scalability / elasticity", "Traffic is bursty during campaigns and sales; ingestion should not block recommendation serving.", "Use an event-driven broker/queue and independently scalable consumers."],
        ["Recommendation relevance", "Poor suggestions waste screen space and erode user trust.", "Use interaction weights, item similarity, seen-item filtering, Precision@5, and coverage monitoring."],
    ], [1.45, 2.65, 2.4])

    doc.add_page_break()
    doc.add_heading("Objective 2 - System Architecture", level=1)
    doc.add_heading("2.1 Architecture Diagram and Component Responsibilities", level=2)
    add_image(doc, diagram_paths[3], "Figure 4: System architecture with ML and non-ML components clearly separated.")
    add_data_table(doc, [
        ["Component", "Type", "Responsibility"],
        ["Storefront", "External", "Emits activity events and renders returned product recommendations."],
        ["API Gateway", "Non-ML", "Single entry point for routing, authentication, and rate limiting."],
        ["Event Tracking Service", "Non-ML", "Validates user activity and publishes events."],
        ["Message Broker / Queue", "Non-ML", "Buffers events and decouples producers from consumers."],
        ["Event Consumer / Updater", "Non-ML", "Processes queued events and updates feature data asynchronously."],
        ["Feature Store", "Data store", "Stores the user-item matrix and item-similarity table."],
        ["Recommendation Service", "ML component", "Ranks candidate products using item-based collaborative filtering."],
        ["Cache", "Data store", "Stores repeated per-user top-N results for fast reads."],
    ], [1.55, 1.0, 3.95])

    doc.add_heading("2.2 Architectural Patterns Selected", level=2)
    add_para(doc, "Pattern 1 - Event-Driven Architecture. User activity arrives continuously and can spike during campaigns. Producers publish events and consumers react independently, which decouples ingestion from model updates and protects the recommendation read path from write bursts. In the prototype, POST /track places events into a queue and a background consumer updates the feature store.")
    add_para(doc, "Pattern 2 - API Gateway. Clients call one stable gateway endpoint. The gateway centralises authentication, rate limiting, and routing, while the internal recommendation service remains focused on ranking. In the prototype, external calls go to port 8000 and are forwarded internally to port 8001.")
    add_callout(doc, "Together, the two patterns implement the desired shape of the system: asynchronous writes for freshness and elasticity, plus fast controlled reads for the storefront experience.", fill="E8F4F2")

    doc.add_heading("2.3 Implementation Summary", level=2)
    add_data_table(doc, [
        ["File", "Purpose", "Pattern / concern demonstrated"],
        ["recommender_engine.py", "Maintains the user-item matrix, computes item similarity, ranks unseen products, reports offline Precision@5.", "ML feature management and recommendation logic"],
        ["recommendation_api.py", "Internal FastAPI app with POST /track, GET /rank, event queue, and background consumer.", "Event-Driven Architecture"],
        ["api_gateway.py", "External FastAPI app with GET /recommend and POST /activity; validates token and routes requests.", "API Gateway"],
        ["demo_requests.py", "Sends events through the gateway and captures the recommendation response.", "Application evidence"],
        ["report_evidence.py", "Generates metric JSON, console output text, and recommendation plot.", "Evaluation and screenshots / figures"],
        ["final_submission/GXXX.ipynb", "Notebook version of the implementation and evidence cells for assignment submission.", "Implementation notebook"],
    ], [1.45, 3.05, 2.0])
    add_para(doc, "Run commands used for the local prototype:")
    add_code_block(doc, """
        python -m pip install -r requirements.txt
        python -m uvicorn recommendation_api:app --host 127.0.0.1 --port 8001
        python -m uvicorn api_gateway:app --host 127.0.0.1 --port 8000
        python demo_requests.py
        python report_evidence.py
    """)

    doc.add_heading("2.4 Application Evidence and Screenshots", level=2)
    add_para(doc, "The following figures come from the live local run. They include real browser-rendered Swagger UI screenshots plus request/response and console evidence. The gateway accepted activity events, routed them to the internal service, and returned recommendations through the single public API.")
    add_image(doc, gateway_swagger, "Figure 5: API Gateway Swagger UI showing the public /health, /recommend, and /activity endpoints.")
    add_image(doc, service_swagger, "Figure 6: Internal recommendation/event service Swagger UI showing /track, /rank, /stats, and /health.")
    add_image(doc, api_evidence, "Figure 7: Live API evidence from gateway request/response output.")
    add_image(doc, consumer_log, "Figure 8: Event consumer log showing asynchronous event processing.")
    add_image(doc, recommendation_plot, "Figure 9: Executed recommendation evidence: matrix view and top-5 products for user u7.")
    if sample_output:
        add_para(doc, "Captured console output from the executed run:")
        add_code_block(doc, sample_output)

    doc.add_heading("2.5 Reading the Result", level=2)
    add_para(doc, "The demonstration sends three fresh actions for user u7 through the gateway. The response marks the activity calls as event-driven and served by the API gateway. After the consumer processes the events, GET /recommend?user_id=u7&k=5 returns items P28, P25, P21, P16, and P40, ranked by collaborative-filtering score. Already-seen items are filtered out, so the output is a true recommendation list rather than an echo of prior interactions.")

    doc.add_heading("2.6 Limitations and Future Improvements", level=2)
    add_para(doc, "This is an academic prototype, so it uses a small deterministic dataset and an in-memory queue to keep the run simple. In production, the event queue should be replaced with Kafka, Redis Streams, or a managed broker; feature data should be stored in a persistent feature store; and monitoring should track latency, event lag, Precision@5, catalogue coverage, drift, and online CTR uplift through A/B testing.")

    doc.add_heading("Appendix A - Key Code Excerpts", level=1)
    add_para(doc, "The full runnable code is in the seml-ecommerce-reco folder and the notebook deliverable is final_submission/GXXX.ipynb. The excerpts below show where the selected patterns and recommendation logic are implemented.")
    doc.add_heading("A.1 Event-Driven queue and consumer", level=2)
    add_code_block(doc, """
        EVENT_QUEUE: Queue[dict[str, str]] = Queue()

        def consumer_loop() -> None:
            while not STOP_WORKER.is_set():
                event = EVENT_QUEUE.get(timeout=0.2)
                result = recommender_engine.track_event(
                    event["user_id"], event["item_id"], event["action"]
                )
                print("processed event:", result, flush=True)
                EVENT_QUEUE.task_done()

        @app.post("/track", status_code=202)
        def track(event: ActivityEvent) -> dict[str, Any]:
            recommender_engine.validate_event(event.user_id, event.item_id, event.action)
            EVENT_QUEUE.put(event.model_dump())
            return {"status": "accepted", "pattern": "event-driven"}
    """)
    doc.add_heading("A.2 API Gateway route", level=2)
    add_code_block(doc, """
        @app.get("/recommend")
        async def recommend(user_id: str, k: int = 5,
                            authorization: str | None = Header(default=None)) -> Any:
            require_token(authorization)
            enforce_rate_limit(f"recommend:{user_id}")
            async with httpx.AsyncClient(timeout=5.0) as client:
                response = await client.get(
                    f"{RECO_SERVICE}/rank",
                    params={"user_id": user_id, "k": k},
                )
            payload = await forward_response(response)
            payload["served_by"] = "api-gateway"
            payload["pattern"] = "api-gateway"
            return payload
    """)
    doc.add_heading("A.3 Collaborative-filtering ranking", level=2)
    add_code_block(doc, """
        def rank_items(user_id: str, k: int = 5) -> list[dict[str, Any]]:
            user_row = _MATRIX[USER_INDEX[user_id]].copy()
            scores = user_row @ _ITEM_SIMILARITY
            scores = scores.astype(float, copy=True)
            scores[user_row > 0] = -np.inf
            ordered = np.argsort(np.nan_to_num(scores, neginf=-1e12))[::-1]
            return [
                {"item_id": ITEMS[i], "score": round(float(scores[i]), 3)}
                for i in ordered
                if user_row[i] == 0
            ][:k]
    """)

    doc.add_heading("Task Distribution and References", level=1)
    add_data_table(doc, [
        ["Member / role", "Tasks owned", "Deliverable", "%"],
        ["Member 1 - Requirements Lead", "Domain, problem statement, requirements, measurable goals, and top quality requirements.", "Sections 1.1, 1.2, 1.4", "25"],
        ["Member 2 - GR4ML Modelling Lead", "Business, Analytics Design and Data Preparation views; diagram validation.", "Section 1.3 + diagrams", "25"],
        ["Member 3 - Architecture Lead", "Architecture diagram, ML/non-ML component split, and pattern justification.", "Sections 2.1, 2.2", "25"],
        ["Member 4 - Implementation Lead", "FastAPI services, event queue, gateway routing, evidence generation, screenshots, and code explanation.", "Sections 2.3, 2.4", "25"],
    ], [1.65, 2.95, 1.4, 0.5])

    doc.add_heading("Submission Readiness Checklist", level=2)
    add_data_table(doc, [
        ["Item", "Status"],
        ["PDF final report", "Generated in final_submission/"],
        ["Editable DOCX report", "Generated in final_submission/"],
        ["Implementation notebook", "Generated as final_submission/GXXX.ipynb; rename to actual group number"],
        ["Runnable source code", "Available in seml-ecommerce-reco/"],
        ["Application screenshots", "Gateway Swagger, internal service Swagger, API evidence, terminal log, output plot"],
        ["Group number, names, BITS IDs", "Must be filled by the group before upload"],
    ], [2.3, 4.2])

    doc.add_heading("References", level=2)
    refs = [
        "Nalchigar, S., and Yu, E. (2020). Designing Business Analytics Solutions: A Model-Driven Approach. Business & Information Systems Engineering, 62(1), 61-75.",
        "Nalchigar, S., and Yu, E. (2018). Business-driven data analytics: A conceptual modeling framework. Data & Knowledge Engineering.",
        "Nalchigar, S., Yu, E., Obeidi, Y. et al. (2019). Solution Patterns for Machine Learning. CAiSE 2019.",
        "Sarwar, B., Karypis, G., Konstan, J., and Riedl, J. (2001). Item-based collaborative filtering recommendation algorithms. WWW 2001.",
        "FastAPI documentation and scikit-learn cosine_similarity documentation used for the local prototype implementation.",
    ]
    for idx, ref in enumerate(refs, start=1):
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(ref)
        set_run_font(r, size=11, color=INK)

    add_callout(
        doc,
        "Final submission checklist: fill group details, verify the file name matches the group ID, and upload this DOCX or the exported PDF to Taxila before 30 June 2026, 23:00.",
    )

    doc.core_properties.title = "SEML Assignment I - E-commerce Recommendation"
    doc.core_properties.author = "Group - BITS Pilani WILP"
    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    print(build_docx())

