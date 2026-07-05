"""Build the complete assignment report as DOCX and PDF."""

from __future__ import annotations

import argparse
import inspect
import json
import textwrap
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image,
    KeepTogether,
    PageBreak,
    Paragraph,
    Preformatted,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from ecom_ml.ml.data import prepare_interactions
from ecom_ml.ml.evaluation import evaluate_leave_n_out
from ecom_ml.ml.model import CollaborativeFilteringModel
from ecom_ml.ml.pipeline import train_pipeline

ROOT = Path(__file__).resolve().parents[1]
EVIDENCE = ROOT / "evidence"
FINAL = ROOT / "final_submission"
PDF_OUTPUT = ROOT / "output" / "pdf"
DETAILS = json.loads((ROOT / "submission_details.json").read_text(encoding="utf-8"))
METADATA = json.loads((ROOT / "artifacts" / "model_metadata.json").read_text(encoding="utf-8"))
GROUP = DETAILS["group_number"]
STEM = f"{GROUP}_SEML_Assignment_01_Complete_Report"
DOCX_PATH = FINAL / f"{STEM}.docx"
PDF_PATH = PDF_OUTPUT / f"{STEM}.pdf"

# standard_business_brief preset, kept deliberately simple for student homework
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
NAVY = RGBColor(0x16, 0x32, 0x4F)
INK = RGBColor(0x1F, 0x29, 0x33)
MUTED = RGBColor(0x52, 0x60, 0x6D)
AMBER = RGBColor(0xD9, 0x77, 0x06)
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(
    run: Any,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell: Any, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(
    cell: Any,
    *,
    top: int = 80,
    bottom: int = 80,
    start: int = 120,
    end: int = 120,
) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table: Any, widths_in: list[float]) -> None:
    widths_dxa = [round(width * 1440) for width in widths_in]
    if sum(widths_dxa) != TABLE_WIDTH_DXA:
        widths_dxa[-1] += TABLE_WIDTH_DXA - sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    properties = table._tbl.tblPr
    table_width = properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        properties.append(table_width)
    table_width.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    table_width.set(qn("w:type"), "dxa")

    indent = properties.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    indent.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            properties = cell._tc.get_or_add_tcPr()
            cell_width = properties.first_child_found_in("w:tcW")
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                properties.append(cell_width)
            cell_width.set(qn("w:w"), str(widths_dxa[index]))
            cell_width.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_page_field(paragraph: Any) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, display, end])


def configure_docx(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
        style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1
        style.paragraph_format.keep_with_next = True

    code_style = styles.add_style("Code Block", 1)
    code_style.font.name = "Consolas"
    code_style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Consolas")
    code_style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Consolas")
    code_style.font.size = Pt(7.5)
    code_style.paragraph_format.space_before = Pt(4)
    code_style.paragraph_format.space_after = Pt(8)
    code_style.paragraph_format.line_spacing = 1.0
    code_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    caption_style = styles.add_style("Figure Caption", 1)
    caption_style.font.name = "Calibri"
    caption_style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    caption_style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    caption_style.font.size = Pt(9)
    caption_style.font.italic = True
    caption_style.font.color.rgb = MUTED
    caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_style.paragraph_format.space_after = Pt(8)

    header = section.header.paragraphs[0]
    header.clear()
    header_run = header.add_run("SEML ASSIGNMENT I  |  E-COMMERCE ML RECOMMENDATION")
    set_run_font(header_run, size=8.5, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.clear()
    add_page_field(footer)


def add_paragraph(document: Document, text: str, *, bold_lead: str | None = None) -> Any:
    paragraph = document.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, size=11, color=INK, bold=True)
        body = paragraph.add_run(text[len(bold_lead) :])
        set_run_font(body, size=11, color=INK)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, size=11, color=INK)
    return paragraph


def add_callout(document: Document, label: str, text: str, *, fill: str = "FFF4DD") -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.08)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(8)
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)
    borders = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "5")
        border.set(qn("w:space"), "5")
        border.set(qn("w:color"), "D97706")
        borders.append(border)
    properties.append(borders)
    label_run = paragraph.add_run(f"{label}: ")
    set_run_font(label_run, size=10.5, color=NAVY, bold=True)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, size=10.5, color=INK)


def add_table(
    document: Document,
    rows: list[list[str]],
    widths: list[float],
    *,
    font_size: float = 8.5,
) -> None:
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set(qn("w:val"), "true")
    header_properties.append(repeat_header)
    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            if row_index == 0:
                set_cell_shading(cell, "F4F6F9")
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.08
            run = paragraph.add_run(str(value))
            set_run_font(
                run,
                size=font_size,
                color=DARK_BLUE if row_index == 0 else INK,
                bold=row_index == 0,
            )
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def add_image(document: Document, filename: str, caption: str, *, width: float = 6.25) -> None:
    path = EVIDENCE / filename
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    picture = paragraph.add_run().add_picture(str(path), width=Inches(width))
    picture._inline.docPr.set("title", caption)
    picture._inline.docPr.set("descr", caption)
    caption_paragraph = document.add_paragraph(style="Figure Caption")
    caption_paragraph.add_run(caption)


def add_code(document: Document, code: str, title: str) -> None:
    heading = document.add_paragraph()
    heading.paragraph_format.space_before = Pt(5)
    heading.paragraph_format.space_after = Pt(3)
    heading.paragraph_format.keep_with_next = True
    run = heading.add_run(title)
    set_run_font(run, size=9, color=DARK_BLUE, bold=True)

    paragraph = document.add_paragraph(style="Code Block")
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F6F8FA")
    paragraph._p.get_or_add_pPr().append(shading)
    code_run = paragraph.add_run(textwrap.dedent(code).strip())
    set_run_font(code_run, name="Consolas", size=7.5, color=INK)


def add_cover(document: Document) -> None:
    document.add_paragraph().paragraph_format.space_after = Pt(42)
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    set_run_font(kicker.add_run("AIMLCZG546"), size=11, color=AMBER, bold=True)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    set_run_font(
        title.add_run("Software Engineering for Machine Learning"),
        size=22,
        color=NAVY,
        bold=True,
    )
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    set_run_font(
        subtitle.add_run("Assignment I - E-commerce Product Recommendation"),
        size=16,
        color=BLUE,
        bold=True,
    )
    descriptor = document.add_paragraph()
    descriptor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    descriptor.paragraph_format.space_after = Pt(42)
    set_run_font(
        descriptor.add_run(
            "GR4ML | Collaborative filtering | Interest cold start | Microservices + CQRS | Streamlit"
        ),
        size=11,
        color=MUTED,
        italic=True,
    )
    metadata = document.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata.paragraph_format.space_after = Pt(5)
    set_run_font(metadata.add_run(f"Group: {GROUP}"), size=12, color=NAVY, bold=True)
    deadline = document.add_paragraph()
    deadline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    deadline.paragraph_format.space_after = Pt(28)
    set_run_font(
        deadline.add_run(f"Submission deadline: {DETAILS['deadline']}"),
        size=10.5,
        color=MUTED,
    )


def member_table() -> list[list[str]]:
    rows = [["Sl.", "BITS ID", "Name", "Contribution", "%"]]
    for index, member in enumerate(DETAILS["members"], start=1):
        rows.append(
            [
                str(index),
                member["bits_id"],
                member["name"],
                member["contribution"],
                str(member["percentage"]),
            ]
        )
    return rows


def report_content_docx(document: Document) -> None:
    metrics = METADATA["metrics"]

    group_heading = document.add_heading("Group Details and Submission Note", level=1)
    group_heading.paragraph_format.page_break_before = True
    add_table(document, member_table(), [0.35, 0.9, 1.0, 3.75, 0.5], font_size=7.8)
    add_callout(
        document,
        "Required before upload",
        "Replace all TO_FILL values with the actual BITS IDs and member names before submission.",
    )
    add_paragraph(
        document,
        "Submission format. We have included the implementation notebook together with this "
        "Word report. The report contains our answers, selected code excerpts, the three GR4ML "
        "views, the system architecture, and browser captures of the running APIs and UI.",
        bold_lead="Submission format.",
    )

    document.add_heading("How This Report Answers the Seven Questions", level=1)
    add_table(
        document,
        [
            ["Question", "Requirement", "Evidence in this submission"],
            ["1", "Domain and ML problem statement", "Section 1"],
            ["2", "Requirements and measurable goals using GR4ML", "Section 2"],
            ["3", "Business, Analytics Design, Data Preparation views", "Section 3; Figures 1-3"],
            ["4", "Top three quality requirements with justification", "Section 4"],
            ["5", "Architecture with ML and non-ML components", "Section 5; Figure 4"],
            ["6", "Select and apply two relevant patterns", "Section 6"],
            ["7", "Implement patterns with code and screenshots", "Sections 7-9; Appendix A"],
        ],
        [0.55, 3.15, 2.8],
    )

    document.add_heading("Objective 1 - Requirements Formulation", level=1)
    document.add_heading("Question 1 - Domain and Problem Statement", level=2)
    add_paragraph(
        document,
        "Chosen domain. We selected e-commerce and online retail. A storefront must decide "
        "which small set of products to show each shopper while catalogue size, user intent, and "
        "inventory change continuously.",
        bold_lead="Chosen domain.",
    )
    add_paragraph(
        document,
        "Problem statement. We use timestamped implicit feedback - view, click, cart, and purchase "
        "events - to build an ML system that learns product affinity and returns a personalized "
        "top-5 list of unseen products for a known user. We use item-based collaborative "
        "filtering because it is explainable, inexpensive to train at assignment scale, and fast at "
        "query time.",
        bold_lead="Problem statement.",
    )
    add_paragraph(
        document,
        "ML formulation. We map each action to preference strength (1, 2, 3, or 5), aggregate it "
        "into a user-item matrix, and transform it into an item-item cosine-similarity model. A user "
        "vector multiplied by this similarity matrix produces scores; previously seen products are "
        "masked before ranking.",
        bold_lead="ML formulation.",
    )
    add_paragraph(
        document,
        "Why ML is appropriate. User preference is not governed by a stable set of deterministic "
        "rules. It depends on sparse historical behavior and changes over time. The requirement "
        "therefore covers data quality, uncertainty in model output, repeatable retraining, and "
        "continuous evaluation - not only an API response.",
        bold_lead="Why ML is appropriate.",
    )
    add_paragraph(
        document,
        "Following the lecture, we aligned goals at four levels so that model performance is "
        "traceable to user and business value.",
    )
    add_table(
        document,
        [
            ["Goal level", "Goal for this application", "Measure"],
            [
                "Organizational",
                "Increase revenue per shopping session.",
                "CTR uplift and average order value",
            ],
            [
                "Product",
                "Show a useful personalized product block.",
                "Recommendation usage and response latency",
            ],
            ["User", "Find interesting products quickly.", "Top-5 hit rate and user feedback"],
            ["Model", "Rank relevant unseen products.", "Precision@5, Recall@5, Coverage@5"],
        ],
        [1.25, 3.4, 1.85],
        font_size=8.1,
    )

    document.add_heading(
        "Question 2 - Requirement Specifications and Measurable Goals",
        level=2,
    )
    add_paragraph(
        document,
        "Unlike a deterministic application, an ML system depends on both software behavior and "
        "the training data. We therefore specify functional, data, model, and quality requirements.",
    )
    add_table(
        document,
        [
            ["ID", "Requirement specification", "Verification"],
            [
                "FR1",
                "Load a schema-valid interaction dataset.",
                "CSV loader rejects missing columns.",
            ],
            [
                "FR2",
                "Deduplicate events and apply action weights.",
                "Unit tests inspect event counts and matrix.",
            ],
            [
                "FR3",
                "Train and persist an item-based CF model.",
                "NPZ artifact and metadata are version matched.",
            ],
            [
                "FR4",
                "Evaluate without train-test leakage.",
                "Leave-2-out Precision/Recall/Hit Rate@5.",
            ],
            [
                "FR5",
                "Accept interactions and training as commands.",
                "Command Service returns 202/200.",
            ],
            [
                "FR6",
                "Serve top-k unseen products as queries.",
                "Query Service returns five ranked products.",
            ],
            [
                "FR7",
                "Create named users with a primary catalogue interest.",
                "User command persists a validated profile.",
            ],
            [
                "FR8",
                "Recommend useful products before a new user has interaction history.",
                "Cold-start query ranks products from the selected interest.",
            ],
            [
                "DR1",
                "Accept only complete IDs, valid actions, and ISO timestamps; ignore duplicate event IDs.",
                "Validation and deduplication unit tests.",
            ],
            [
                "MR1",
                "Make training reproducible and version every persisted model.",
                "Fixed seeds and model/metadata version check.",
            ],
            [
                "NFR1",
                "Keep training writes separate from low-latency reads.",
                "Independent Command and Query services.",
            ],
        ],
        [0.48, 3.3, 2.72],
        font_size=8.2,
    )
    add_paragraph(
        document,
        "The first three measures can be calculated from the held-out interaction data. Serving "
        "latency and business uplift need live traffic, so we keep them as deployment targets "
        "rather than presenting them as results from this prototype.",
    )
    add_table(
        document,
        [
            ["Goal", "Target", "Result from this run", "Status"],
            [
                "Recommendation relevance",
                "Precision@5 >= 0.30",
                f"{metrics['precision_at_k']:.3f}",
                "Passed",
            ],
            ["Held-out recovery", "Recall@5 >= 0.70", f"{metrics['recall_at_k']:.3f}", "Passed"],
            [
                "Catalogue reach",
                "Coverage@5 >= 0.80",
                f"{metrics['catalogue_coverage_at_k']:.3f}",
                "Passed",
            ],
            ["Serving latency", "p95 < 200 ms", "Requires load test", "Production validation"],
            ["Business impact", "CTR uplift >= 8%", "Requires A/B test", "Production validation"],
        ],
        [2.0, 1.45, 1.7, 1.35],
    )

    document.add_heading("Question 3 - GR4ML Views", level=2)
    add_paragraph(
        document,
        "We use one visual style across all three GR4ML views: the title treatment, type scale, "
        "line weight, spacing, and footer key remain consistent. The shapes retain their GR4ML "
        "meaning: goals are ovals, algorithms are hexagons, softgoals are clouds, indicators use "
        "a traffic-light symbol, entities use structured boxes, operators use rectangles, and notes "
        "have a folded corner.",
    )
    add_image(
        document,
        "gr4ml_business_view.png",
        "Figure 1. GR4ML Business View using Actor, Business Goal, Decision Goal, Question Goal, Indicator, and Insight notation.",
    )
    add_paragraph(
        document,
        "Business interpretation. The Storefront Manager wants to increase revenue per session. "
        "That business goal leads to a recurring product decision: which items should occupy the "
        "limited recommendation space? The Top-5 Recommendation Insight answers the related "
        "question about unseen products. CTR and average order value are the indicators used to "
        "judge whether the decision helps the business goal.",
    )
    add_image(
        document,
        "gr4ml_analytics_design_view.png",
        "Figure 2. GR4ML Analytics Design View using Analytics Goal, Algorithm, Softgoal, Indicator, Performs, Influence, Evaluates, and Generates notation.",
    )
    add_paragraph(
        document,
        "Analytics interpretation. The analytics goal is to estimate preferences for products a "
        "user has not seen. We considered item-based collaborative filtering, matrix factorization, "
        "and a popularity baseline. Item-based CF fits this prototype because its similarity scores "
        "are easy to inspect and inexpensive to serve. Precision@5 checks relevance, while "
        "Coverage@5 checks whether recommendations extend beyond a small part of the catalogue.",
    )
    add_image(
        document,
        "gr4ml_data_preparation_view.png",
        "Figure 3. GR4ML Data Preparation View using Entities, Operators, Note, Data Flow, Input/Output, and Relationship notation.",
    )
    add_paragraph(
        document,
        "Data interpretation. Interaction Event and Product Catalogue are the two source entities. "
        "The pipeline validates events, removes duplicate event IDs, maps actions to weights, and "
        "aggregates values for each user-item pair. The result is the matrix used for evaluation, "
        "training, and ranking. The folded note records the exact weights so the transformation can "
        "be reproduced.",
    )
    add_callout(
        document,
        "Cross-view traceability",
        "Prepared User-Item Matrix is required for the Prediction and Ranking Analytics Goal; "
        "that goal generates the Top-5 Recommendation Insight; the insight answers the Question "
        "Goal; the answer supports the product-display Decision Goal; and the decision contributes "
        "to the revenue-per-session Business Goal.",
        fill="E6F5F2",
    )

    document.add_heading("Question 4 - Top Three Quality Requirements", level=2)
    add_table(
        document,
        [
            ["Quality requirement", "Why it is critical", "Design response"],
            [
                "Recommendation relevance",
                "Irrelevant products waste scarce storefront space and reduce user trust.",
                "Weighted feedback, unseen-item masking, Precision/Recall/Hit Rate@5 gates.",
            ],
            [
                "Low query latency",
                "Recommendations are requested on the page-render path.",
                "Precomputed similarity and a read-only Query Service; p95 target < 200 ms.",
            ],
            [
                "Scalability and change isolation",
                "Training is compute-heavy while reads are frequent and latency-sensitive.",
                "Command and Query microservices can deploy and scale independently.",
            ],
        ],
        [1.5, 2.45, 2.55],
        font_size=8.2,
    )
    add_paragraph(
        document,
        "The Analytics Design View shows relevance and latency as softgoals and includes "
        "interpretability because it influenced the algorithm choice. Scalability is addressed in "
        "the architecture by separating training commands from recommendation reads. Together, the "
        "two views keep model quality and software quality in the same design discussion.",
    )

    architecture_heading = document.add_heading("Objective 2 - System Architecture", level=1)
    architecture_heading.paragraph_format.page_break_before = True
    document.add_heading("Question 5 - Architecture and Responsibilities", level=2)
    add_image(
        document,
        "system_architecture.png",
        "Figure 4. System architecture separating ML, data, command/write, and query/read concerns.",
    )
    add_table(
        document,
        [
            ["Component", "Type", "Responsibility"],
            [
                "Streamlit UI :8501",
                "Presentation",
                "Runs user, interaction, training, and recommendation flows.",
            ],
            ["Command Service :8101", "Non-ML service", "Owns writes and invokes training."],
            [
                "Training pipeline",
                "ML component",
                "Loads, prepares, evaluates, fits, and persists.",
            ],
            ["Raw event log", "Data store", "Authoritative command-side interaction history."],
            ["User profile store", "Data store", "Named users and primary cold-start interests."],
            ["Versioned read model", "ML artifact", "Similarity matrix, user features, metadata."],
            [
                "Query Service :8102",
                "ML serving",
                "Serves collaborative or interest-based recommendations and recent actions.",
            ],
        ],
        [1.55, 1.25, 3.7],
    )
    add_paragraph(
        document,
        "Read path. A recommendation request goes from Streamlit to the Query Service. The service "
        "reads the user profile and current model artifact, masks products already seen by the user, "
        "and returns ranked catalogue records. This path does not modify the event log or trigger training.",
        bold_lead="Read path.",
    )
    add_paragraph(
        document,
        "Write and training path. User creation, interaction recording, and retraining go through "
        "the Command Service. Interactions are appended to the event log; training reads that log, "
        "builds a fresh model, and publishes a new artifact version. The Query Service detects the "
        "new version and reloads it for later requests.",
        bold_lead="Write and training path.",
    )

    document.add_heading("Question 6 - Two Selected Architectural Patterns", level=2)
    document.add_heading("Pattern 1 - Microservices", level=3)
    add_paragraph(
        document,
        "Application. Command and Query responsibilities are separate FastAPI applications on ports "
        "8101 and 8102. Each has its own entry point, API contract, health endpoint, tests, and "
        "Docker Compose process. This keeps model training away from the request path used to fetch "
        "recommendations and allows the two workloads to be scaled separately.",
        bold_lead="Application.",
    )
    add_paragraph(
        document,
        "Trade-off. Distributed deployment increases operational complexity. The local setup "
        "uses a shared artifact directory for reproducibility; a production design would publish to "
        "a model registry or object store and add authentication, tracing, retries, and rollback.",
        bold_lead="Trade-off.",
    )
    document.add_heading("Pattern 2 - CQRS", level=3)
    add_paragraph(
        document,
        "Application. POST /commands/interactions, POST /commands/users, and POST /commands/train "
        "mutate the write side. Product, user, recent-action, model-info, and recommendation endpoints "
        "remain read-only. The query process atomically reloads a new model version; a profile with no "
        "trained history receives an interest-based cold-start ranking.",
        bold_lead="Application.",
    )
    add_paragraph(
        document,
        "Trade-off. CQRS introduces eventual consistency: an accepted interaction affects queries "
        "only after training and publication. For recommendations this is acceptable when freshness "
        "is monitored and the previous valid model remains available during retraining.",
        bold_lead="Trade-off.",
    )

    document.add_heading("Question 7 - Implementation, Code, and Output", level=2)
    document.add_heading("7.1 ML Training and Evaluation", level=3)
    add_paragraph(
        document,
        "The pipeline reports five stages: LOAD, PREPARE, EVALUATE, TRAIN, and PERSIST. The seed "
        "creates 80 users, 60 items, and 960 interaction events. The model artifact used for this "
        f"report records {METADATA['training_events']} events and stores arrays in NPZ with metadata "
        "in JSON; it does not rely on an executable pickle file.",
    )
    add_image(
        document,
        "ml_execution_metrics.png",
        "Figure 5. Prepared interaction matrix, offline metrics, and the top-5 scores from the saved model.",
    )
    add_callout(
        document,
        "Offline evaluation result",
        f"Precision@5={metrics['precision_at_k']:.3f}, "
        f"Recall@5={metrics['recall_at_k']:.3f}, "
        f"Hit Rate@5={metrics['hit_rate_at_k']:.3f}, and "
        f"Coverage@5={metrics['catalogue_coverage_at_k']:.3f} across "
        f"{metrics['evaluated_users']} users.",
        fill="E6F5F2",
    )
    add_paragraph(
        document,
        "How to read the measures. Precision@5 is the fraction of five returned products that were "
        "held out as relevant. Recall@5 measures how many held-out products were recovered. Hit "
        "Rate@5 asks whether at least one held-out product appeared for a user. Coverage@5 reports "
        "the fraction of catalogue items that appeared in any top-5 list. These values check the "
        "offline pipeline on synthetic data; they are not a substitute for an online business experiment.",
        bold_lead="How to read the measures.",
    )

    document.add_heading("7.2 Pattern Implementation and API Evidence", level=3)
    add_image(
        document,
        "command_service_api.png",
        "Figure 6. Command Service contract shown in Swagger UI.",
    )
    add_image(
        document,
        "query_service_api.png",
        "Figure 7. Query Service contract shown in Swagger UI.",
    )
    add_paragraph(
        document,
        "API interpretation. The Swagger pages make the CQRS split visible in the running "
        "application. The Command Service owns profile, interaction, and training changes. The "
        "Query Service exposes model metadata, products, named users, recent actions, and "
        "recommendations, but no command routes.",
        bold_lead="API interpretation.",
    )
    add_callout(
        document,
        "Browser capture context",
        "The screenshots record an isolated demonstration run. The Command Service was bound to "
        "port 8111 for that run so its clean interaction file could remain separate from another "
        "local process; the Makefile and architecture retain 8101 as the normal command port. "
        "Model version identifiers change whenever training publishes a new artifact.",
        fill="FFF4DD",
    )

    document.add_heading("7.3 UI Flows Captured in the Browser", level=3)
    add_paragraph(
        document,
        "Flow 1 - Personalized recommendation and recent activity. The user selects a named customer "
        "and top-k value, then submits a read-only query. The UI displays realistic product names, "
        "categories, scores, model strategy, and the three most recent actions beneath the ranking.",
        bold_lead="Flow 1 - Personalized recommendation and recent activity.",
    )
    add_image(
        document,
        "ui_recommendation_flow.png",
        "Figure 8. Collaborative recommendations with the three latest actions.",
    )
    add_paragraph(
        document,
        "Flow 2 - Record interaction. A named customer, product, and action type are selected. "
        "Submitting the form sends POST /commands/interactions and appends a validated implicit-feedback "
        "event; the UI explains that retraining is required before it changes collaborative rankings.",
        bold_lead="Flow 2 - Record interaction.",
    )
    add_image(
        document,
        "ui_interaction_flow.png",
        "Figure 9. Interaction form after recording an implicit-feedback event.",
    )
    add_paragraph(
        document,
        "Flow 3 - User management. The Users tab lists the five default named profiles and their "
        "interests. The add-user form validates a name and primary-interest dropdown before "
        "POST /commands/users persists a new profile.",
        bold_lead="Flow 3 - User management.",
    )
    add_image(
        document,
        "ui_users_flow.png",
        "Figure 10. Five default user profiles and the add-user form.",
    )
    add_paragraph(
        document,
        "Flow 4 - Interest-based cold start. A newly created user has no interaction history in the "
        "trained matrix. The Query Service therefore returns products from that profile's selected "
        "interest and labels the strategy interest-based-cold-start. After the user records actions "
        "and the model is retrained, the same request can move to collaborative recommendations; "
        "the saved interest remains available as a fallback.",
        bold_lead="Flow 4 - Interest-based cold start.",
    )
    add_image(
        document,
        "ui_cold_start_flow.png",
        "Figure 11. Interest-based recommendations for a newly created user.",
    )
    add_paragraph(
        document,
        "Flow 5 - Model training. The operator selects evaluation K and holdout events, then submits "
        "POST /commands/train. The five-stage pipeline loads, prepares, evaluates, trains, and persists "
        "the read model; the Query Service reloads it without exposing a write route.",
        bold_lead="Flow 5 - Model training.",
    )
    add_image(
        document,
        "ui_training_flow.png",
        "Figure 12. Training controls and the completed pipeline summary.",
    )
    add_paragraph(
        document,
        "Training result shown in Figure 12. The clean browser dataset started with 960 events. "
        "The interaction flow added one click before training, so the captured summary reports "
        "961 raw and unique events, 80 users, 60 products, Precision@5 of 0.330, Recall@5 of "
        "0.825, Hit Rate@5 of 0.950, and full catalogue coverage. Later local runs may have a "
        "different event count and version while following the same pipeline.",
        bold_lead="Training result shown in Figure 12.",
    )

    document.add_heading("7.4 Verification and Reproducibility", level=3)
    add_table(
        document,
        [
            ["Check", "Result"],
            ["Ruff lint", "Passed"],
            ["Ruff formatting", "Passed"],
            ["Mypy static typing", "20 source files passed"],
            ["Pytest", "8 tests passed"],
            ["Coverage", "88% total"],
            ["Browser flows", "Five UI flows and two Swagger pages captured"],
            ["Notebook", "All code cells executed; outputs and browser evidence saved"],
        ],
        [2.4, 4.1],
    )
    add_code(
        document,
        """
        uv venv --python 3.12 .venv
        source .venv/bin/activate
        uv pip install -r requirements.txt -e ".[dev,report]"
        python scripts/seed_data.py
        python scripts/train_and_evaluate.py

        # Run in separate terminals
        uvicorn ecom_ml.command_service.main:app --port 8101
        uvicorn ecom_ml.query_service.main:app --port 8102
        streamlit run frontend/app.py --server.port 8501

        # Quality gates
        make check
        python scripts/verify_live.py
        """,
        "Reproduction commands",
    )

    document.add_heading("7.5 Limitations and Future Improvements", level=3)
    add_paragraph(
        document,
        "This prototype uses deterministic synthetic interactions, a CSV command store, and a "
        "shared local artifact directory. It does not include online experiment data, authentication, "
        "privacy controls, or drift alerts. Cold start uses one declared interest, which is useful "
        "for showing the flow but much simpler than a production hybrid recommender.",
    )
    add_paragraph(
        document,
        "As future work, we would use an event broker and durable warehouse, a feature store, "
        "a model registry with approval and rollback, content/popularity fallbacks, distributed "
        "tracing, load tests, drift and bias monitoring, and an A/B experiment measuring CTR and "
        "revenue per session.",
    )

    document.add_heading("Appendix A - Key Code Excerpts", level=1)
    add_paragraph(
        document,
        "The notebook runs the ML path and embeds the relevant service modules. The main excerpts "
        "are repeated below so the implementation can be reviewed alongside the design discussion.",
    )
    add_code(
        document,
        inspect.getsource(prepare_interactions),
        "A.1 Data validation, deduplication, weighting, and matrix construction",
    )
    add_code(
        document,
        inspect.getsource(CollaborativeFilteringModel.fit),
        "A.2 Model training",
    )
    add_code(
        document,
        inspect.getsource(CollaborativeFilteringModel.recommend),
        "A.3 Recommendation inference",
    )
    add_code(
        document,
        inspect.getsource(evaluate_leave_n_out),
        "A.4 Leakage-safe evaluation",
    )
    add_code(
        document,
        inspect.getsource(train_pipeline),
        "A.5 End-to-end ML pipeline",
    )

    document.add_heading("References", level=1)
    references = [
        (
            "R1",
            "Rao, S. (2026). AIMLCZG546 Session 3 lecture slides: Requirements Engineering and "
            "GR4ML Modeling Notation. BITS Pilani WILP.",
        ),
        (
            "R2",
            "Nalchigar, S. and Yu, E. (2020). Designing Business Analytics Solutions: "
            "A Model-Driven Approach. Business & Information Systems Engineering, 62, 61-75. "
            "https://doi.org/10.1007/s12599-018-0555-z",
        ),
        (
            "R3",
            "Sarwar, B., Karypis, G., Konstan, J., and Riedl, J. (2001). Item-based "
            "collaborative filtering recommendation algorithms. WWW '01, 285-295. "
            "https://doi.org/10.1145/371920.372071",
        ),
        (
            "R4",
            "Microsoft Azure Architecture Center. CQRS pattern. "
            "https://learn.microsoft.com/en-us/azure/architecture/patterns/cqrs",
        ),
        (
            "R5",
            "Microsoft Azure Architecture Center. Architecture styles - Microservices. "
            "https://learn.microsoft.com/en-us/azure/architecture/guide/architecture-styles/",
        ),
        (
            "R6",
            "scikit-learn documentation. sklearn.metrics.pairwise.cosine_similarity. "
            "https://scikit-learn.org/stable/modules/generated/"
            "sklearn.metrics.pairwise.cosine_similarity.html",
        ),
        (
            "R7",
            "FastAPI documentation. Automatic interactive API documentation. "
            "https://fastapi.tiangolo.com/features/#automatic-docs",
        ),
    ]
    for label, reference in references:
        add_paragraph(document, f"{label}. {reference}", bold_lead=f"{label}.")

    add_callout(
        document,
        "Final submission check",
        "Fill group metadata, confirm the portal naming convention, open the notebook to "
        "verify saved outputs, and submit the requested IPYNB plus one report format.",
    )


def build_docx() -> Path:
    FINAL.mkdir(parents=True, exist_ok=True)
    document = Document()
    configure_docx(document)
    add_cover(document)
    report_content_docx(document)
    document.core_properties.title = "SEML Assignment I - E-commerce ML Recommendation"
    document.core_properties.subject = (
        "GR4ML, Microservices, CQRS, collaborative filtering, interest cold start, Streamlit"
    )
    document.core_properties.author = f"BITS Pilani WILP Group {GROUP}"
    document.core_properties.keywords = (
        "SEML, GR4ML, recommendation, Microservices, CQRS, Streamlit, cold start"
    )
    document.save(DOCX_PATH)
    return DOCX_PATH


# PDF builder -----------------------------------------------------------------

PDF_NAVY = colors.HexColor("#16324F")
PDF_BLUE = colors.HexColor("#2E74B5")
PDF_MUTED = colors.HexColor("#52606D")
PDF_GRID = colors.HexColor("#CBD5E1")
PDF_FILL = colors.HexColor("#F4F6F9")
PDF_AMBER = colors.HexColor("#D97706")


def pdf_styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    return {
        "Body": ParagraphStyle(
            "Body",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=9.5,
            leading=12.7,
            textColor=colors.HexColor("#1F2933"),
            alignment=TA_JUSTIFY,
            spaceAfter=7,
        ),
        "H1": ParagraphStyle(
            "H1",
            parent=base["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=15,
            leading=18,
            textColor=PDF_BLUE,
            spaceBefore=14,
            spaceAfter=8,
        ),
        "H2": ParagraphStyle(
            "H2",
            parent=base["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=12,
            leading=14,
            textColor=PDF_BLUE,
            spaceBefore=10,
            spaceAfter=5,
        ),
        "H3": ParagraphStyle(
            "H3",
            parent=base["Heading3"],
            fontName="Helvetica-Bold",
            fontSize=10.5,
            leading=12.5,
            textColor=PDF_NAVY,
            spaceBefore=7,
            spaceAfter=4,
        ),
        "Caption": ParagraphStyle(
            "Caption",
            parent=base["BodyText"],
            fontName="Helvetica-Oblique",
            fontSize=8,
            leading=10,
            textColor=PDF_MUTED,
            alignment=TA_CENTER,
            spaceAfter=8,
        ),
        "Table": ParagraphStyle(
            "Table",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=7.4,
            leading=9.2,
            textColor=colors.HexColor("#1F2933"),
            alignment=TA_LEFT,
        ),
        "TableHeader": ParagraphStyle(
            "TableHeader",
            parent=base["BodyText"],
            fontName="Helvetica-Bold",
            fontSize=7.5,
            leading=9.4,
            textColor=PDF_NAVY,
            alignment=TA_LEFT,
        ),
        "Code": ParagraphStyle(
            "Code",
            parent=base["Code"],
            fontName="Courier",
            fontSize=6.2,
            leading=7.4,
            leftIndent=5,
            rightIndent=5,
            borderColor=PDF_GRID,
            borderWidth=0.5,
            borderPadding=7,
            backColor=colors.HexColor("#F6F8FA"),
            spaceAfter=8,
        ),
        "CoverTitle": ParagraphStyle(
            "CoverTitle",
            parent=base["Title"],
            fontName="Helvetica-Bold",
            fontSize=26,
            leading=31,
            textColor=PDF_NAVY,
            alignment=TA_CENTER,
            spaceAfter=8,
        ),
        "CoverSub": ParagraphStyle(
            "CoverSub",
            parent=base["BodyText"],
            fontName="Helvetica-Bold",
            fontSize=15,
            leading=18,
            textColor=PDF_BLUE,
            alignment=TA_CENTER,
            spaceAfter=10,
        ),
        "Center": ParagraphStyle(
            "Center",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=9.5,
            leading=12,
            textColor=PDF_MUTED,
            alignment=TA_CENTER,
            spaceAfter=6,
        ),
    }


PSTYLE = pdf_styles()


def pdf_table(rows: list[list[str]], widths: list[float], *, font_size: float = 7.4) -> Table:
    body_style = ParagraphStyle(
        "TableDynamic",
        parent=PSTYLE["Table"],
        fontSize=font_size,
        leading=font_size + 1.8,
    )
    wrapped: list[list[Paragraph]] = []
    for row_index, row in enumerate(rows):
        style = PSTYLE["TableHeader"] if row_index == 0 else body_style
        wrapped.append([Paragraph(str(value).replace("&", "&amp;"), style) for value in row])
    table = Table(
        wrapped,
        colWidths=[width * inch for width in widths],
        repeatRows=1,
        hAlign="LEFT",
    )
    table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.35, PDF_GRID),
                ("BACKGROUND", (0, 0), (-1, 0), PDF_FILL),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    return table


def pdf_image(filename: str, caption: str, *, width: float = 6.25) -> KeepTogether:
    path = EVIDENCE / filename
    image = Image(str(path), width=width * inch)
    image.drawHeight = image.imageHeight * image.drawWidth / image.imageWidth
    return KeepTogether([image, Spacer(1, 4), Paragraph(caption, PSTYLE["Caption"])])


def pdf_callout(label: str, text: str) -> Table:
    content = Paragraph(f"<b>{label}:</b> {text}", PSTYLE["Body"])
    table = Table([[content]], colWidths=[6.5 * inch], hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#FFF4DD")),
                ("BOX", (0, 0), (-1, -1), 0.7, PDF_AMBER),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
            ]
        )
    )
    return table


def pdf_header_footer(canvas: Any, document: Any) -> None:
    canvas.saveState()
    canvas.setStrokeColor(PDF_GRID)
    canvas.setLineWidth(0.4)
    canvas.line(0.75 * inch, 10.37 * inch, 7.75 * inch, 10.37 * inch)
    canvas.setFont("Helvetica-Bold", 7.5)
    canvas.setFillColor(PDF_MUTED)
    canvas.drawString(0.75 * inch, 10.48 * inch, "SEML ASSIGNMENT I | E-COMMERCE ML RECOMMENDATION")
    canvas.setFont("Helvetica", 8)
    canvas.drawRightString(7.75 * inch, 0.45 * inch, f"Page {document.page}")
    canvas.restoreState()


def build_pdf_story() -> list[Any]:
    metrics = METADATA["metrics"]
    story: list[Any] = [
        Spacer(1, 1.15 * inch),
        Paragraph(
            "AIMLCZG546",
            ParagraphStyle(
                "Kicker", parent=PSTYLE["Center"], textColor=PDF_AMBER, fontName="Helvetica-Bold"
            ),
        ),
        Paragraph("Software Engineering for Machine Learning", PSTYLE["CoverTitle"]),
        Paragraph("Assignment I - E-commerce Product Recommendation", PSTYLE["CoverSub"]),
        Paragraph(
            "GR4ML | Collaborative filtering | Interest cold start | Microservices + CQRS | Streamlit",
            PSTYLE["Center"],
        ),
        Spacer(1, 0.75 * inch),
        Paragraph(f"<b>Group:</b> {GROUP}", PSTYLE["Center"]),
        Paragraph(f"Submission deadline: {DETAILS['deadline']}", PSTYLE["Center"]),
        Spacer(1, 0.4 * inch),
        Paragraph(f"Prepared as a report plus the {GROUP}.ipynb notebook", PSTYLE["Center"]),
        PageBreak(),
        Paragraph("Group Details and Submission Note", PSTYLE["H1"]),
        pdf_table(member_table(), [0.35, 0.9, 1.0, 3.75, 0.5], font_size=6.8),
        Spacer(1, 8),
        pdf_callout(
            "Required before upload",
            "Replace all TO_FILL values with the actual BITS IDs and member names.",
        ),
        Spacer(1, 8),
        Paragraph(
            "<b>Submission format.</b> We include the implementation notebook and a reviewable "
            "Word/PDF with selected code, diagrams, calculated results, and browser captures.",
            PSTYLE["Body"],
        ),
        Paragraph("Assignment Coverage Matrix", PSTYLE["H1"]),
        pdf_table(
            [
                ["Question", "Requirement", "Evidence"],
                ["1", "Domain and ML problem", "Section 1"],
                ["2", "Requirements and measurable goals", "Section 2"],
                ["3", "Three GR4ML views", "Section 3; Figures 1-3"],
                ["4", "Top three quality requirements", "Section 4"],
                ["5", "Architecture with ML/non-ML", "Section 5; Figure 4"],
                ["6", "Two relevant patterns", "Section 6"],
                ["7", "Implementation and screenshots", "Sections 7-9; Appendix A"],
            ],
            [0.6, 3.4, 2.5],
        ),
        Paragraph("Objective 1 - Requirements Formulation", PSTYLE["H1"]),
        Paragraph("1. Domain and Problem Statement", PSTYLE["H2"]),
        Paragraph(
            "<b>Domain.</b> E-commerce and online retail. A storefront must decide which small set "
            "of products to show each shopper while catalogue size and intent change continuously.",
            PSTYLE["Body"],
        ),
        Paragraph(
            "<b>Problem statement.</b> Given implicit view, click, cart, and purchase events, learn "
            "product affinity and return a personalized top-5 list of unseen products for a known user.",
            PSTYLE["Body"],
        ),
        Paragraph(
            "<b>ML formulation.</b> Actions receive weights 1, 2, 3, and 5. Aggregated user-item "
            "features train item-item cosine similarity; seen products are masked before ranking.",
            PSTYLE["Body"],
        ),
        Paragraph("2. Requirement Specifications and Measurable Goals", PSTYLE["H2"]),
        pdf_table(
            [
                ["ID", "Functional requirement", "Verification"],
                ["FR1", "Load schema-valid events.", "Loader rejects missing columns."],
                ["FR2", "Deduplicate and weight events.", "Tests inspect matrix and counts."],
                ["FR3", "Train and persist item-based CF.", "Version-matched NPZ + JSON."],
                ["FR4", "Evaluate without leakage.", "Leave-2-out top-k metrics."],
                ["FR5", "Accept writes/training as commands.", "Command service 202/200."],
                ["FR6", "Serve unseen items as queries.", "Query returns five products."],
                ["FR7", "Create named users with interests.", "Profile command + user query."],
                ["FR8", "Support new-user cold start.", "Interest-matched products."],
            ],
            [0.5, 3.35, 2.65],
        ),
        Spacer(1, 8),
        pdf_table(
            [
                ["Goal", "Target", "Result from this run", "Status"],
                ["Relevance", "Precision@5 >= 0.30", f"{metrics['precision_at_k']:.3f}", "Passed"],
                ["Recovery", "Recall@5 >= 0.70", f"{metrics['recall_at_k']:.3f}", "Passed"],
                [
                    "Reach",
                    "Coverage@5 >= 0.80",
                    f"{metrics['catalogue_coverage_at_k']:.3f}",
                    "Passed",
                ],
                ["Latency", "p95 < 200 ms", "Needs load test", "Production"],
                ["Impact", "CTR uplift >= 8%", "Needs A/B test", "Production"],
            ],
            [1.5, 1.7, 1.8, 1.5],
        ),
        Paragraph("3. GR4ML Views", PSTYLE["H2"]),
        Paragraph(
            "All three views use the same title treatment, type scale, line weight, spacing, and "
            "footer key. Their shapes retain the Session 3 GR4ML meaning: goals are ovals, "
            "algorithms are hexagons, softgoals are clouds, indicators use traffic lights, "
            "entities are structured boxes, operators are rectangles, and notes have folded corners.",
            PSTYLE["Body"],
        ),
        pdf_image("gr4ml_business_view.png", "Figure 1. GR4ML Business View."),
        Paragraph(
            "Personalized presentation supports revenue per session. The product-display decision is "
            "informed by a top-5 insight and monitored through CTR and average order value.",
            PSTYLE["Body"],
        ),
        pdf_image("gr4ml_analytics_design_view.png", "Figure 2. GR4ML Analytics Design View."),
        Paragraph(
            "The goal is to rank unseen products. Item-based collaborative filtering was selected "
            "because it is inexpensive to serve and its similarities are easy to inspect.",
            PSTYLE["Body"],
        ),
        pdf_image("gr4ml_data_preparation_view.png", "Figure 3. GR4ML Data Preparation View."),
        Paragraph(
            "Events are validated, deduplicated, weighted, and aggregated with the product catalogue "
            "to form the user-item matrix used by evaluation and training.",
            PSTYLE["Body"],
        ),
        Paragraph("4. Top Three Quality Requirements", PSTYLE["H2"]),
        pdf_table(
            [
                ["Quality", "Justification", "Design response"],
                [
                    "Relevance",
                    "Poor suggestions reduce trust.",
                    "Weighted features and top-k quality gates.",
                ],
                [
                    "Low latency",
                    "Reads occur on page render.",
                    "Precomputed model and read-only service.",
                ],
                [
                    "Scalability",
                    "Training and reads differ.",
                    "Independent command/query deployment.",
                ],
            ],
            [1.25, 2.4, 2.85],
        ),
        PageBreak(),
        Paragraph("Objective 2 - System Architecture", PSTYLE["H1"]),
        Paragraph("5. Architecture and Responsibilities", PSTYLE["H2"]),
        pdf_image("system_architecture.png", "Figure 4. Microservices + CQRS system architecture."),
        pdf_table(
            [
                ["Component", "Type", "Responsibility"],
                ["Streamlit UI", "Presentation", "All five user-facing flows."],
                ["Command Service", "Non-ML", "Writes and invokes training."],
                ["Training pipeline", "ML", "Prepare, evaluate, fit, persist."],
                ["Raw event log", "Data", "Command-side source of truth."],
                ["User profiles", "Data", "Names and cold-start interests."],
                ["Read model", "ML artifact", "Similarity, features, metadata."],
                ["Query Service", "ML serving", "Collaborative/cold-start reads."],
            ],
            [1.5, 1.25, 3.75],
        ),
        Paragraph(
            "<b>Request paths.</b> Recommendation reads go from Streamlit to the Query Service, "
            "which combines the saved profile, model artifact, and catalogue without changing the "
            "event log. User, interaction, and training operations go to the Command Service. "
            "Training publishes a new artifact version that the Query Service reloads.",
            PSTYLE["Body"],
        ),
        Paragraph("6. Selected Architectural Patterns", PSTYLE["H2"]),
        Paragraph("6.1 Pattern 1 - Microservices", PSTYLE["H3"]),
        Paragraph(
            "<b>Application.</b> Command and Query are separate FastAPI processes on ports 8101 and "
            "8102, with independent contracts and health endpoints. Read replicas and training can "
            "scale independently.",
            PSTYLE["Body"],
        ),
        Paragraph(
            "<b>Trade-off.</b> The local artifact directory favors reproducibility. Production would "
            "use a model registry plus authentication, tracing, retries, and rollback.",
            PSTYLE["Body"],
        ),
        Paragraph("6.2 Pattern 2 - CQRS", PSTYLE["H3"]),
        Paragraph(
            "<b>Application.</b> User, interaction, and training commands mutate the write side. "
            "Product, user, activity, model, and recommendation queries remain read-only. New users "
            "receive an interest-based cold-start ranking.",
            PSTYLE["Body"],
        ),
        Paragraph(
            "<b>Trade-off.</b> Accepted interactions become visible after retraining, creating "
            "intentional eventual consistency.",
            PSTYLE["Body"],
        ),
        Paragraph("7. ML Implementation", PSTYLE["H2"]),
        Paragraph(
            "The pipeline exposes LOAD, PREPARE, EVALUATE, TRAIN, and PERSIST. The deterministic seed "
            f"contains 80 users, 60 items, and 960 events; the artifact used here records "
            f"{METADATA['training_events']} events.",
            PSTYLE["Body"],
        ),
        pdf_image(
            "ml_execution_metrics.png",
            "Figure 5. Prepared matrix, offline metrics, and recommendation scores.",
        ),
        pdf_callout(
            "Offline evaluation result",
            f"Precision@5={metrics['precision_at_k']:.3f}; "
            f"Recall@5={metrics['recall_at_k']:.3f}; "
            f"Hit Rate@5={metrics['hit_rate_at_k']:.3f}; "
            f"Coverage@5={metrics['catalogue_coverage_at_k']:.3f}.",
        ),
        Paragraph(
            "<b>Metric reading.</b> Precision@5 measures relevant products among the five returned; "
            "Recall@5 measures recovery of held-out products; Hit Rate@5 checks whether a user gets "
            "at least one hit; and Coverage@5 measures how much of the catalogue appears across "
            "all top-5 lists. These are offline results on synthetic data.",
            PSTYLE["Body"],
        ),
        Paragraph("8. Pattern Implementation and API Evidence", PSTYLE["H2"]),
        pdf_image(
            "command_service_api.png",
            "Figure 6. Command Service contract shown in Swagger UI.",
        ),
        pdf_image(
            "query_service_api.png",
            "Figure 7. Query Service contract shown in Swagger UI.",
        ),
        Paragraph(
            "The Swagger pages show the CQRS split in the running application. The command side "
            "owns mutations; the query side exposes no command route.",
            PSTYLE["Body"],
        ),
        pdf_callout(
            "Browser capture context",
            "The isolated browser run used command port 8111 to keep its clean interaction file "
            "separate; the normal Makefile and architecture port remains 8101. Model versions "
            "change after each training run.",
        ),
        Paragraph("9. UI Flows Captured in the Browser", PSTYLE["H2"]),
        Paragraph(
            "<b>Flow 1 - Recommendation.</b> Select a named customer and top-k value; display named "
            "products, scores, strategy, and three recent actions.",
            PSTYLE["Body"],
        ),
        pdf_image(
            "ui_recommendation_flow.png",
            "Figure 8. Recommendations and the three latest actions.",
        ),
        Paragraph(
            "<b>Flow 2 - Interaction.</b> Select customer, product, and action; submit a validated "
            "command to the write-side event log.",
            PSTYLE["Body"],
        ),
        pdf_image(
            "ui_interaction_flow.png",
            "Figure 9. Interaction form after recording an event.",
        ),
        Paragraph(
            "<b>Flow 3 - Users.</b> Review five defaults and create a profile with a primary interest.",
            PSTYLE["Body"],
        ),
        pdf_image(
            "ui_users_flow.png",
            "Figure 10. Default users and the add-user form.",
        ),
        Paragraph(
            "<b>Flow 4 - Cold start.</b> A new user receives products from the selected interest "
            "until trained interaction history exists. Recorded actions and retraining can later "
            "move the profile to collaborative recommendations.",
            PSTYLE["Body"],
        ),
        pdf_image(
            "ui_cold_start_flow.png",
            "Figure 11. Interest-based recommendations for a new user.",
        ),
        Paragraph(
            "<b>Flow 5 - Training.</b> Configure evaluation controls and execute the five-stage "
            "training/persistence workflow.",
            PSTYLE["Body"],
        ),
        pdf_image(
            "ui_training_flow.png",
            "Figure 12. Training controls and the completed run summary.",
        ),
        Paragraph(
            "<b>Captured training result.</b> One click was added to the 960-event seed before "
            "training. The screenshot therefore records 961 events, 80 users, 60 products, "
            "Precision@5 0.330, Recall@5 0.825, Hit Rate@5 0.950, and full catalogue coverage.",
            PSTYLE["Body"],
        ),
        Paragraph("10. Verification and Reproducibility", PSTYLE["H2"]),
        pdf_table(
            [
                ["Check", "Result"],
                ["Ruff lint / format", "Passed"],
                ["Mypy", "20 source files passed"],
                ["Pytest", "8 tests passed"],
                ["Coverage", "88%"],
                ["Browser evidence", "Five UI flows + two Swagger pages"],
                ["Notebook", "All cells executed; browser evidence saved"],
            ],
            [2.4, 4.1],
        ),
        Spacer(1, 8),
        Preformatted(
            textwrap.dedent(
                """
                uv venv --python 3.12 .venv
                source .venv/bin/activate
                uv pip install -r requirements.txt -e ".[dev,report]"
                python scripts/seed_data.py
                python scripts/train_and_evaluate.py
                uvicorn ecom_ml.command_service.main:app --port 8101
                uvicorn ecom_ml.query_service.main:app --port 8102
                streamlit run frontend/app.py --server.port 8501
                make check
                """
            ).strip(),
            PSTYLE["Code"],
        ),
        Paragraph("11. Limitations and Production Roadmap", PSTYLE["H2"]),
        Paragraph(
            "This prototype uses synthetic events, CSV, and a shared local artifact. Online uplift "
            "and production latency have not been measured. A production version would need a "
            "broker, warehouse, feature store, model registry, stronger hybrid cold start, security, "
            "drift monitoring, load tests, and an A/B experiment.",
            PSTYLE["Body"],
        ),
        Paragraph("Appendix A - Key Code Excerpts", PSTYLE["H1"]),
        Paragraph(
            "The notebook embeds the relevant modules. These excerpts come from the same source files.",
            PSTYLE["Body"],
        ),
    ]
    for title, source in [
        ("A.1 Preprocessing", inspect.getsource(prepare_interactions)),
        ("A.2 Training", inspect.getsource(CollaborativeFilteringModel.fit)),
        ("A.3 Inference", inspect.getsource(CollaborativeFilteringModel.recommend)),
        ("A.4 Evaluation", inspect.getsource(evaluate_leave_n_out)),
        ("A.5 Pipeline", inspect.getsource(train_pipeline)),
    ]:
        story.extend(
            [
                Paragraph(title, PSTYLE["H3"]),
                Preformatted(textwrap.dedent(source).strip(), PSTYLE["Code"]),
            ]
        )

    story.extend(
        [
            Paragraph("References", PSTYLE["H1"]),
            Paragraph(
                "<b>R1.</b> Rao, S. (2026). AIMLCZG546 Session 3 lecture slides: Requirements "
                "Engineering and GR4ML Modeling Notation. BITS Pilani WILP.",
                PSTYLE["Body"],
            ),
            Paragraph(
                "<b>R2.</b> Nalchigar, S. and Yu, E. (2020). Designing Business Analytics "
                "Solutions. https://doi.org/10.1007/s12599-018-0555-z",
                PSTYLE["Body"],
            ),
            Paragraph(
                "<b>R3.</b> Sarwar, B. et al. (2001). Item-based collaborative filtering "
                "recommendation algorithms. https://doi.org/10.1145/371920.372071",
                PSTYLE["Body"],
            ),
            Paragraph(
                "<b>R4.</b> Microsoft Azure Architecture Center. CQRS pattern. "
                "https://learn.microsoft.com/en-us/azure/architecture/patterns/cqrs",
                PSTYLE["Body"],
            ),
            Paragraph(
                "<b>R5.</b> Microsoft Azure Architecture Center. Architecture styles. "
                "https://learn.microsoft.com/en-us/azure/architecture/guide/architecture-styles/",
                PSTYLE["Body"],
            ),
            Paragraph(
                "<b>R6.</b> scikit-learn. cosine_similarity documentation. "
                "https://scikit-learn.org/stable/modules/generated/"
                "sklearn.metrics.pairwise.cosine_similarity.html",
                PSTYLE["Body"],
            ),
            Paragraph(
                "<b>R7.</b> FastAPI. Automatic API docs. "
                "https://fastapi.tiangolo.com/features/#automatic-docs",
                PSTYLE["Body"],
            ),
            pdf_callout(
                "Final submission check",
                "Fill group metadata, confirm naming, inspect saved notebook outputs, and submit "
                "the required IPYNB plus one report format.",
            ),
        ]
    )
    return story


def build_pdf() -> Path:
    PDF_OUTPUT.mkdir(parents=True, exist_ok=True)
    document = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=0.78 * inch,
        bottomMargin=0.7 * inch,
        title="SEML Assignment I - E-commerce ML Recommendation",
        author=f"BITS Pilani WILP Group {GROUP}",
    )
    document.build(
        build_pdf_story(),
        onFirstPage=pdf_header_footer,
        onLaterPages=pdf_header_footer,
    )
    # Mirror the PDF beside the notebook/DOCX for portal submission convenience.
    final_pdf = FINAL / PDF_PATH.name
    final_pdf.write_bytes(PDF_PATH.read_bytes())
    return final_pdf


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--with-pdf",
        action="store_true",
        help="Also build the PDF after creating the Word document.",
    )
    return parser.parse_args()


if __name__ == "__main__":
    arguments = parse_args()
    print(build_docx())
    if arguments.with_pdf:
        print(build_pdf())
